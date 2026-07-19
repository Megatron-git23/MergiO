from PyPDF2 import PdfMerger
from docx import Document

def merge_pdfs():
          file1 = input("Enter the path of PDF file 1: ")
          file2 = input("Enter the path of PDF file 2: ")
          output = input("Enter the output file name: ")

          merger = PdfMerger()

          merger.append(file1)
          merger.append(file2)

          merger.write(output)
          merger.close()

          print(f"PDFs merged successfully into '{output}.pdf'")


def merge_docx():
    file1 = input("Enter the path of Word file 1: ")
    file2 = input("Enter the path of Word file 2: ")
    output = input("Enter the output file name: ")

    merged = Document()

    if merged.paragraphs:
        p = merged.paragraphs[0]._element
        p.getparent().remove(p)

    for file in [file1, file2]:
        document = Document(file)

        for element in document.element.body:
            merged.element.body.append(element)

    merged.save(output)

    print(f"Word files merged successfully into '{output}.docx'")



choice=int(input("Enter your choice:\n Enter '1' for pdf merger: \n Enter '2' for docx merger:"))

if choice==1:
    merge_pdfs()
else:
    merge_docx()

